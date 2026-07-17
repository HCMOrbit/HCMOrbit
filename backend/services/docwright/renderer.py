"""Render a Docwright document to .docx and .pdf.

.docx:  python-docx, real heading styles, real tables, page numbers, title page.
.pdf:   markdown → HTML → PDF via xhtml2pdf (pure Python, no system deps).

Both renderers consume the same section dict produced by
`services.docwright.generator.generate_document`.
"""
from __future__ import annotations

import io
import re
from datetime import datetime, timezone
from html import escape as _hesc
from typing import Iterable

import markdown as _md
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from xhtml2pdf import pisa

from services.docwright.generator import SECTION_KEYS, SECTION_KEY_TO_LABEL

# Brand palette — matches HCMOrbit teal/dark tokens used in the UI.
BRAND_INK = RGBColor(0x0A, 0x16, 0x28)
BRAND_TEAL = RGBColor(0x0D, 0x93, 0x73)
OPEN_ITEM_HEX = "#F5B731"  # amber, for PDF highlight

# ── shared helpers ──────────────────────────────────────────────────────────

# Very-lightweight markdown table detector — `_render_markdown_docx` splits
# a markdown blob into "table blocks" and "everything else" so python-docx
# can render real tables instead of dumping pipes as text.
_TABLE_LINE_RE = re.compile(r"^\s*\|.*\|\s*$")
_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")


def _split_markdown_blocks(md: str) -> list[tuple[str, str]]:
    """Return [(block_type, block_text)] where block_type ∈ {"table","text"}."""
    lines = (md or "").splitlines()
    blocks: list[tuple[str, str]] = []
    buf: list[str] = []
    buf_kind = "text"

    def flush():
        if buf:
            blocks.append((buf_kind, "\n".join(buf)))
        buf.clear()

    i = 0
    while i < len(lines):
        ln = lines[i]
        # A table block starts when we see a header row followed by a separator
        if _TABLE_LINE_RE.match(ln) and i + 1 < len(lines) and _TABLE_SEP_RE.match(lines[i + 1]):
            # Flush the pending text block
            if buf_kind != "table":
                flush(); buf_kind = "table"
            # Consume the table until the pipes stop
            while i < len(lines) and _TABLE_LINE_RE.match(lines[i]):
                buf.append(lines[i]); i += 1
            flush(); buf_kind = "text"
            continue
        if buf_kind != "text":
            flush(); buf_kind = "text"
        buf.append(ln); i += 1
    flush()
    return blocks


def _parse_markdown_table(block: str) -> tuple[list[str], list[list[str]]]:
    """Parse a markdown table block into (headers, rows)."""
    lines = [ln.strip() for ln in block.splitlines() if ln.strip()]
    if len(lines) < 2:
        return [], []
    def _cells(ln: str) -> list[str]:
        raw = ln.strip().strip("|")
        return [c.strip() for c in raw.split("|")]
    header = _cells(lines[0])
    # Skip the separator line (line 1)
    rows = [_cells(ln) for ln in lines[2:]]
    # Pad short rows to header width so python-docx doesn't crash
    for r in rows:
        while len(r) < len(header):
            r.append("")
    return header, rows


# ── .docx renderer ─────────────────────────────────────────────────────────

def _add_page_number_field(paragraph):
    """Insert a Word PAGE field so the footer shows the page number."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)


def _apply_body_font(doc: Document) -> None:
    # Base body font: Calibri 11pt (Word default; safest cross-platform).
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _render_bold_inline(run_container, text: str) -> None:
    """Render `**bold**` runs in `text` onto a python-docx paragraph.

    OPEN ITEM: markers get amber-highlighted so the client-facing docx is
    obvious about gaps.
    """
    # Also highlight OPEN ITEM: prefix even if not bold
    # First split on **...** while preserving delimiter groups
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = False
        content = part
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            bold = True; content = part[2:-2]
        is_open_item = "OPEN ITEM:" in content
        run = run_container.add_run(content)
        run.bold = bold or is_open_item
        if is_open_item:
            # Amber highlight for OPEN ITEM markers
            run.font.color.rgb = RGBColor(0x8A, 0x61, 0x00)


def _render_markdown_text(doc: Document, text: str) -> None:
    """Render a markdown text block onto the doc. Handles headings, lists,
    bold, and OPEN ITEM highlights. Deliberately basic — clients want a clean
    document, not full markdown fidelity."""
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph()
            continue
        # H3 → level 2, H2 → level 2, H1 → level 1
        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3); continue
        if line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2); continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1); continue
        # Bullet lists
        if line.lstrip().startswith(("- ", "* ", "• ")):
            p = doc.add_paragraph(style="List Bullet")
            _render_bold_inline(p, line.lstrip()[2:].lstrip())
            continue
        # Numbered lists (1. Foo)
        m = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _render_bold_inline(p, m.group(1))
            continue
        p = doc.add_paragraph()
        _render_bold_inline(p, line)


def _render_markdown_table(doc: Document, block: str) -> None:
    header, rows = _parse_markdown_table(block)
    if not header:
        return
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.style = "Light Grid Accent 1"
    # Header row
    for j, h in enumerate(header):
        cell = tbl.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    # Data rows
    for i, r in enumerate(rows, start=1):
        for j in range(len(header)):
            cell = tbl.rows[i].cells[j]
            cell.text = ""
            _render_bold_inline(cell.paragraphs[0], r[j])
    doc.add_paragraph()


def _render_section(doc: Document, block_md: str) -> None:
    for kind, block in _split_markdown_blocks(block_md):
        if kind == "table":
            _render_markdown_table(doc, block)
        else:
            _render_markdown_text(doc, block)


def render_docx(*, meta: dict, sections: dict) -> bytes:
    """Return a .docx blob for the given metadata + sections dict."""
    doc = Document()
    _apply_body_font(doc)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("\n\n\n\n")
    r.font.size = Pt(1)  # spacer

    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hdr.add_run(meta.get("doc_type") or "Configuration Design Document")
    run.bold = True; run.font.size = Pt(28); run.font.color.rgb = BRAND_INK

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"{meta.get('client_name','')}  ·  {meta.get('module','')}")
    run.font.size = Pt(14); run.font.color.rgb = BRAND_TEAL

    phase = doc.add_paragraph()
    phase.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = phase.add_run(f"Phase: {meta.get('phase','')}")
    run.font.size = Pt(12)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(
        datetime.now(timezone.utc).strftime("%d %B %Y")
    )
    run.font.size = Pt(11); run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_page_break()

    # Footer with page numbers
    footer_p = doc.sections[0].footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number_field(footer_p)

    # Body sections
    for key, label in SECTION_KEYS:
        doc.add_heading(label, level=1)
        content = (sections or {}).get(key) or f"**OPEN ITEM:** {label} — no content provided."
        _render_section(doc, content)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF renderer ───────────────────────────────────────────────────────────

_PDF_STYLES = f"""
<style>
  @page {{ size: A4; margin: 22mm 18mm 20mm 18mm; @frame footer {{ -pdf-frame-content: footer_content; bottom: 8mm; margin-left: 18mm; margin-right: 18mm; height: 8mm; }} }}
  body {{ font-family: Helvetica, Arial, sans-serif; font-size: 10.5pt; color: #0A1628; line-height: 1.5; }}
  .title-page {{ text-align: center; padding-top: 60mm; }}
  .title-page h1 {{ font-size: 28pt; color: #0A1628; margin: 0 0 8mm 0; }}
  .title-page .sub {{ font-size: 14pt; color: #0D9373; margin: 0 0 3mm 0; }}
  .title-page .phase {{ font-size: 12pt; color: #0A1628; margin: 0 0 3mm 0; }}
  .title-page .date {{ font-size: 10pt; color: #64748B; }}
  h1 {{ font-size: 16pt; color: #0A1628; border-bottom: 1px solid #E2E8F0; padding-bottom: 3mm; margin: 8mm 0 4mm 0; }}
  h2 {{ font-size: 13pt; color: #0A1628; margin: 6mm 0 3mm 0; }}
  h3 {{ font-size: 11pt; color: #0A1628; margin: 4mm 0 2mm 0; }}
  p  {{ margin: 0 0 3mm 0; }}
  ul, ol {{ margin: 0 0 3mm 6mm; padding: 0; }}
  li {{ margin: 0 0 1.5mm 0; }}
  table {{ border-collapse: collapse; width: 100%; margin: 3mm 0; }}
  th, td {{ border: 1px solid #CBD5E1; padding: 6px 8px; font-size: 9.5pt; vertical-align: top; text-align: left; }}
  th {{ background: #F1F5F9; }}
  .open-item {{ background: #FEF3C7; color: #8A6100; padding: 1px 4px; border-radius: 3px; font-weight: bold; }}
  .footer {{ font-size: 8.5pt; color: #94A3B8; text-align: center; }}
</style>
"""


def _highlight_open_items(html: str) -> str:
    """Wrap `OPEN ITEM: <phrase>` spans in an amber-highlighted class.
    Handles both `<strong>OPEN ITEM:</strong> ...` (markdown ** …**) and
    plain `OPEN ITEM: ...`."""
    # markdown → <strong>OPEN ITEM:</strong> — merge with the phrase after
    html = re.sub(
        r"<strong>\s*OPEN ITEM:\s*</strong>\s*([^<\n]*)",
        r'<span class="open-item">OPEN ITEM: \1</span>',
        html, flags=re.IGNORECASE,
    )
    html = re.sub(
        r"(?<!>)OPEN ITEM:\s*([^<\n]*)",
        r'<span class="open-item">OPEN ITEM: \1</span>',
        html, flags=re.IGNORECASE,
    )
    return html


def render_pdf(*, meta: dict, sections: dict) -> bytes:
    """Return a PDF blob for the given metadata + sections dict."""
    body_parts: list[str] = []
    for key, label in SECTION_KEYS:
        md = (sections or {}).get(key) or f"**OPEN ITEM:** {label} — no content provided."
        html_body = _md.markdown(md, extensions=["tables", "fenced_code"])
        html_body = _highlight_open_items(html_body)
        body_parts.append(f"<h1>{_hesc(label)}</h1>\n{html_body}")

    title_page = f"""
    <div class="title-page">
      <h1>{_hesc(meta.get('doc_type') or 'Configuration Design Document')}</h1>
      <div class="sub">{_hesc(meta.get('client_name',''))} · {_hesc(meta.get('module',''))}</div>
      <div class="phase">Phase: {_hesc(meta.get('phase',''))}</div>
      <div class="date">{datetime.now(timezone.utc).strftime('%d %B %Y')}</div>
    </div>
    <pdf:nextpage />
    """

    footer = (
        '<div id="footer_content" class="footer">'
        f'{_hesc(meta.get("client_name",""))} · {_hesc(meta.get("doc_type",""))} · '
        'Page <pdf:pagenumber /> of <pdf:pagecount />'
        '</div>'
    )

    html = f"""<!doctype html>
<html><head><meta charset="utf-8">{_PDF_STYLES}</head>
<body>{title_page}{''.join(body_parts)}{footer}</body></html>"""

    buf = io.BytesIO()
    pisa_status = pisa.CreatePDF(src=html, dest=buf, encoding="utf-8")
    if pisa_status.err:
        # xhtml2pdf logs the error internally; surface a clean 500 message.
        raise RuntimeError("PDF rendering failed — check backend logs for xhtml2pdf details.")
    return buf.getvalue()
