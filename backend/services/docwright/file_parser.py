"""Parse uploaded consultant notes into plain text.

Supports .docx (via python-docx), .txt, and .md. .doc (legacy binary) is
NOT supported — the frontend restricts the accept attribute to prevent it.
"""
from __future__ import annotations

import io
from typing import Optional

from docx import Document  # python-docx


class FileParseError(RuntimeError):
    """Raised on unreadable / unsupported / oversized uploads."""


MAX_UPLOAD_BYTES = 5 * 1024 * 1024  # 5 MB — plenty for consultant notes


def parse_upload(filename: str, blob: bytes) -> str:
    """Return plain text extracted from `blob`. Raises FileParseError on failure."""
    if not blob:
        raise FileParseError("Uploaded file is empty.")
    if len(blob) > MAX_UPLOAD_BYTES:
        raise FileParseError(
            f"File exceeds 5 MB limit ({len(blob) // 1024} KB uploaded)."
        )
    name = (filename or "").lower().strip()
    if name.endswith(".docx"):
        return _parse_docx(blob)
    if name.endswith(".txt") or name.endswith(".md"):
        try:
            return blob.decode("utf-8", errors="replace")
        except Exception as e:  # pragma: no cover
            raise FileParseError(f"Could not decode text file: {e}") from e
    raise FileParseError(
        "Unsupported file type. Upload a .docx, .txt, or .md file."
    )


def _parse_docx(blob: bytes) -> str:
    try:
        doc = Document(io.BytesIO(blob))
    except Exception as e:
        raise FileParseError(f"Could not read .docx file: {e}") from e
    parts: list[str] = []
    # Paragraphs, in order
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    # Tables — flatten cells row-by-row so table content isn't dropped
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise FileParseError("The .docx file appears to be empty.")
    return text
