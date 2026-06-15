"""Parse uploaded .docx Knowledge Base articles.

Convention (per HCMOrbit KB upload spec):
- The first table in the document carries the 11-row metadata block:
    Document ID, Title, Module, Sub-Module, Artifact Type, Difficulty,
    Estimated Read Time, Tags, Category, Author, Platform
- Every body element AFTER that first table is the article body and is
  converted to Markdown.

Authoring guideline for callouts (notes / tips / warnings):
- Use Word's built-in `Quote` (or `Intense Quote`) paragraph style and
  start the text with one of these labels followed by a colon:
      Note:      → teal info box
      Tip:       → green lightbulb box
      Warning:   → amber warning box
      Important: → amber warning box
- These paragraphs become Markdown blockquotes (`> ...`) on upload and
  the frontend renders them as styled callout boxes — the icon + colour
  conveys the type, so the literal "Note:" word is stripped from the
  displayed text on the article page.
- Do NOT use coloured text, highlighting, or bordered single-cell tables
  for callouts — they won't be detected.

Returns a dict ready for the review screen — nothing is persisted here.
"""
from __future__ import annotations

import io
import re
from typing import Optional

from docx import Document
from docx.oxml.ns import qn


# --- canonical maps used to flag unknown values ----------------------------

# Module label -> canonical category slug. Keys are case-insensitive
# (normalised to lowercase before lookup). Includes common aliases so admins
# can drop in their existing .docx headers without manual remapping.
KNOWN_MODULE_TO_CATEGORY = {
    # Core HCM
    "core hcm": "core-hcm",
    "core-hcm": "core-hcm",
    "hcm": "core-hcm",
    "human capital management": "core-hcm",

    # Talent Acquisition
    "talent acquisition": "talent-acquisition",
    "recruiting": "talent-acquisition",
    "recruitment": "talent-acquisition",
    "hiring": "talent-acquisition",
    "candidate experience": "talent-acquisition",

    # Talent Management
    "talent management": "talent-management",
    "talent": "talent-management",
    "performance": "talent-management",
    "performance management": "talent-management",
    "succession": "talent-management",
    "succession planning": "talent-management",
    "goals": "talent-management",
    "careers": "talent-management",
    "calibration": "talent-management",

    # Compensation & Benefits
    "compensation & benefits": "compensation-benefits",
    "compensation and benefits": "compensation-benefits",
    "comp & benefits": "compensation-benefits",
    "compensation": "compensation-benefits",
    "comp": "compensation-benefits",
    "benefits": "compensation-benefits",
    "total rewards": "compensation-benefits",
    "advanced compensation": "compensation-benefits",

    # Workforce Management
    "workforce management": "workforce-management",
    "wfm": "workforce-management",
    "time tracking": "workforce-management",
    "time & absence": "workforce-management",
    "time and absence": "workforce-management",
    "absence": "workforce-management",
    "absence management": "workforce-management",
    "scheduling": "workforce-management",

    # Payroll
    "payroll": "payroll",
    "us payroll": "payroll",
    "uk payroll": "payroll",
    "global payroll": "payroll",
    "payroll control center": "payroll",
    "pcc": "payroll",
    "ecp": "payroll",
    "peci": "payroll",

    # Learning & Employee Experience
    "learning & employee experience": "learning-employee-experience",
    "learning and employee experience": "learning-employee-experience",
    "learning": "learning-employee-experience",
    "workday learning": "learning-employee-experience",
    "lms": "learning-employee-experience",
    "employee experience": "learning-employee-experience",
    "employee voice": "learning-employee-experience",
    "journeys": "learning-employee-experience",
    "onboarding": "learning-employee-experience",

    # Workforce Planning & Analytics
    "workforce planning & analytics": "workforce-planning-analytics",
    "workforce planning and analytics": "workforce-planning-analytics",
    "workforce planning": "workforce-planning-analytics",
    "headcount planning": "workforce-planning-analytics",
    "people analytics": "workforce-planning-analytics",
    "workforce analytics": "workforce-planning-analytics",
    "skills cloud": "workforce-planning-analytics",

    # Finance & Accounting
    "finance & accounting": "finance-accounting",
    "finance and accounting": "finance-accounting",
    "finance": "finance-accounting",
    "financials": "finance-accounting",
    "accounting": "finance-accounting",
    "accounting center": "finance-accounting",
    "general ledger": "finance-accounting",
    "intercompany": "finance-accounting",
    "banking": "finance-accounting",

    # Procurement & Spend Management
    "procurement & spend management": "procurement-spend-management",
    "procurement and spend management": "procurement-spend-management",
    "procurement": "procurement-spend-management",
    "spend management": "procurement-spend-management",
    "supplier accounts": "procurement-spend-management",
    "supplier management": "procurement-spend-management",
    "expenses": "procurement-spend-management",
    "sourcing": "procurement-spend-management",
    "strategic sourcing": "procurement-spend-management",

    # Projects & Professional Services
    "projects & professional services": "projects-professional-services",
    "projects and professional services": "projects-professional-services",
    "projects": "projects-professional-services",
    "professional services": "projects-professional-services",
    "psa": "projects-professional-services",
    "resource management": "projects-professional-services",
    "project billing": "projects-professional-services",

    # Planning
    "planning": "planning",
    "adaptive planning": "planning",
    "adaptive": "planning",
    "financial planning": "planning",
    "sales planning": "planning",
    "operational planning": "planning",

    # Analytics & Reporting
    "analytics & reporting": "analytics-reporting",
    "analytics and reporting": "analytics-reporting",
    "reporting & analytics": "analytics-reporting",
    "reporting and analytics": "analytics-reporting",
    "reporting": "analytics-reporting",
    "analytics": "analytics-reporting",
    "birt": "analytics-reporting",
    "prism": "analytics-reporting",
    "prism analytics": "analytics-reporting",
    "dashboards": "analytics-reporting",

    # Integration & Platform
    "integration & platform": "integration-platform",
    "integration and platform": "integration-platform",
    "integration": "integration-platform",
    "integrations": "integration-platform",
    "platform": "integration-platform",
    "studio": "integration-platform",
    "workday studio": "integration-platform",
    "eib": "integration-platform",
    "core connectors": "integration-platform",
    "extend": "integration-platform",
    "workday extend": "integration-platform",
    "rest api": "integration-platform",
    "soap api": "integration-platform",

    # Security & Compliance
    "security & compliance": "security-compliance",
    "security and compliance": "security-compliance",
    "security & roles": "security-compliance",
    "security and roles": "security-compliance",
    "security": "security-compliance",
    "roles": "security-compliance",
    "compliance": "security-compliance",
    "audit": "security-compliance",
    "controls": "security-compliance",
    "sox": "security-compliance",

    # AI & Automation
    "ai & automation": "ai-automation",
    "ai and automation": "ai-automation",
    "ai": "ai-automation",
    "automation": "ai-automation",
    "illuminate": "ai-automation",
    "workday illuminate": "ai-automation",
    "machine learning": "ai-automation",
    "ml": "ai-automation",
    "agents": "ai-automation",

    # Industry Solutions
    "industry solutions": "industry-solutions",
    "industries": "industry-solutions",
    "healthcare": "industry-solutions",
    "higher education": "industry-solutions",
    "government": "industry-solutions",
    "financial services": "industry-solutions",
    "retail": "industry-solutions",
    "manufacturing": "industry-solutions",
}

# Category-cell label -> existing doc_type enum
KNOWN_CATEGORY_TO_TYPE = {
    "fix guide": "fix_guide",
    "fix": "fix_guide",
    "how-to": "how_to",
    "how to": "how_to",
    "learning bite": "learning_bite",
    "reference": "reference",
    "checklist": "checklist",
}

DIFFICULTY_MAP = {
    "beginner": "beginner",
    "intermediate": "intermediate",
    "advanced": "advanced",
    "expert": "advanced",
}


# --- helpers ---------------------------------------------------------------

def _runs_to_markdown(paragraph) -> str:
    """Walk the runs of a paragraph and inline-format bold/italic/code.

    Any leading/trailing whitespace on a styled run is moved OUTSIDE the
    delimiter pair so the result is valid CommonMark: e.g. a bold run
    "Top-level org: " (trailing space) becomes "**Top-level org:** ",
    not "**Top-level org: **" which CommonMark renders as literal asterisks.
    """
    out = []
    for run in paragraph.runs:
        text = run.text or ""
        if not text:
            continue
        text = text.replace("\\", "\\\\")
        bold = bool(run.bold)
        italic = bool(run.italic)
        if not (bold or italic):
            out.append(text)
            continue
        # Split into [leading_ws, core, trailing_ws]
        m = re.match(r"^(\s*)(.*?)(\s*)$", text, re.DOTALL)
        leading, core, trailing = m.group(1), m.group(2), m.group(3)
        if not core:
            # Pure whitespace run — keep as-is, don't wrap
            out.append(text)
            continue
        if bold and italic:
            wrapped = f"***{core}***"
        elif bold:
            wrapped = f"**{core}**"
        else:
            wrapped = f"*{core}*"
        out.append(f"{leading}{wrapped}{trailing}")
    return "".join(out)


def _list_level(paragraph) -> Optional[int]:
    """Return the numbering indent level (0-based) from w:pPr/w:numPr/w:ilvl,
    or None if the paragraph is not part of a numbered/bulleted list."""
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    ilvl = numPr.find(qn("w:ilvl"))
    try:
        return int(ilvl.get(qn("w:val")))
    except (TypeError, ValueError, AttributeError):
        return 0


def _paragraph_to_markdown(paragraph) -> str:
    style = (paragraph.style.name or "").strip() if paragraph.style else ""
    text = _runs_to_markdown(paragraph).rstrip()
    if not text:
        return ""
    if style == "Title":
        return f"# {text}"
    if style == "Heading 1":
        return f"## {text}"
    if style == "Heading 2":
        return f"### {text}"
    if style == "Heading 3":
        return f"#### {text}"
    if style == "Heading 4":
        return f"##### {text}"
    if "List Bullet" in style or "List Paragraph" in style or "List Number" in style:
        level = _list_level(paragraph) or 0
        indent = "  " * level
        marker = "1." if "List Number" in style else "-"
        return f"{indent}{marker} {text}"
    if style == "Quote" or style == "Intense Quote":
        return f"> {text}"
    if "Code" in style:
        return f"`{text}`"
    return text


def _table_to_markdown(table) -> str:
    rows = []
    for r in table.rows:
        cells = [c.text.replace("\n", " ").strip() for c in r.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if not rows:
        return ""
    header_sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
    return "\n".join([rows[0], header_sep, *rows[1:]])


def _read_metadata_table(table) -> dict[str, str]:
    """First table is `label | value` rows. Returns case-normalised dict."""
    md: dict[str, str] = {}
    for row in table.rows:
        if len(row.cells) < 2:
            continue
        label = row.cells[0].text.strip()
        value = row.cells[1].text.strip()
        if label:
            md[label] = value
    return md


def _split_tags(raw: str) -> list[str]:
    if not raw:
        return []
    parts = re.split(r"[,;\n]", raw)
    return [p.strip().lower() for p in parts if p.strip()][:8]


def _normalise_difficulty(raw: str) -> str:
    if not raw:
        return "intermediate"
    key = raw.strip().lower()
    return DIFFICULTY_MAP.get(key, "intermediate")


def _normalise_read_time(raw: str) -> Optional[str]:
    if not raw:
        return None
    # Normalise the en-dash and the spelled-out "minutes" to a tighter form
    cleaned = raw.replace("—", "-").replace("–", "-").strip()
    cleaned = re.sub(r"\s+minutes?\b", " min", cleaned, flags=re.I)
    return cleaned or None


# --- main entry point ------------------------------------------------------

def parse_kb_docx(file_bytes: bytes) -> dict:
    doc = Document(io.BytesIO(file_bytes))
    if not doc.tables:
        raise ValueError("Document does not contain a metadata table at the top.")
    meta = _read_metadata_table(doc.tables[0])
    meta_table_el = doc.tables[0]._element

    # Walk top-level body elements after the metadata table
    body_lines: list[str] = []
    seen_meta = False
    for child in doc.element.body.iterchildren():
        if child is meta_table_el:
            seen_meta = True
            continue
        if not seen_meta:
            continue
        tag = child.tag.split("}")[-1]
        if tag == "p":
            # Wrap with python-docx Paragraph for style/runs access
            from docx.text.paragraph import Paragraph
            para = Paragraph(child, doc.tables[0]._parent if False else doc)  # noqa
            md_line = _paragraph_to_markdown(para)
            body_lines.append(md_line)
        elif tag == "tbl":
            from docx.table import Table
            tbl = Table(child, doc.tables[0]._parent if False else doc)  # noqa
            body_lines.append("")
            body_lines.append(_table_to_markdown(tbl))
            body_lines.append("")

    # Collapse 3+ blank lines into 2
    body_md = "\n".join(body_lines).strip()
    body_md = re.sub(r"\n{3,}", "\n\n", body_md)

    # --- field mapping ---
    module_raw = meta.get("Module", "")
    category_slug = KNOWN_MODULE_TO_CATEGORY.get(module_raw.strip().lower())
    category_flag = None
    if module_raw and not category_slug:
        category_flag = f"Module '{module_raw}' is not a recognised category — pick or create one."

    cat_raw = meta.get("Category", "")
    doc_type = KNOWN_CATEGORY_TO_TYPE.get(cat_raw.strip().lower())
    type_flag = None
    if cat_raw and not doc_type:
        type_flag = f"Artifact category '{cat_raw}' isn't a known doc type — pick the closest match."

    parsed = {
        "title": meta.get("Title", "").strip(),
        "summary": _build_default_summary(body_md),
        "body": body_md,
        "category_slug": category_slug or "",
        "category_label_raw": module_raw,
        "doc_type": doc_type or "reference",
        "doc_type_label_raw": cat_raw,
        "difficulty": _normalise_difficulty(meta.get("Difficulty", "")),
        "tags": _split_tags(meta.get("Tags", "")),
        "reference_id": meta.get("Document ID", "").strip() or None,
        "sub_module": meta.get("Sub-Module", "").strip() or None,
        "read_time": _normalise_read_time(meta.get("Estimated Read Time", "")),
        "platform": meta.get("Platform", "").strip() or "Workday",
        "author_raw": meta.get("Author", "").strip() or None,
        "artifact_type_raw": meta.get("Artifact Type", "").strip() or None,
        "target_groups": ["aspirant", "practitioner", "employer"],
        "flags": [f for f in (category_flag, type_flag) if f],
        "raw_metadata": meta,
    }
    return parsed


def _build_default_summary(body_md: str) -> str:
    """First non-heading paragraph, truncated to ~280 chars, as a default summary."""
    for line in body_md.splitlines():
        s = line.strip()
        if not s or s.startswith("#") or s.startswith("-") or s.startswith(">"):
            continue
        # Strip simple inline markdown
        plain = re.sub(r"[*_`]", "", s)
        if len(plain) > 280:
            return plain[:277].rstrip() + "..."
        return plain
    return ""
